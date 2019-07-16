"""
    created on june 17, 2019
    @author: Risasi
    pip install python-docx
    
"""

from docx import Document

doc = Document()
doc.save('test.docx')
doc.add_paragraph('This is my first paragraph!')
doc.add_paragraph('This is my Second  paragraph!')
doc.add_table(rows = 2, cols = 4)
doc.save('test.docx')
